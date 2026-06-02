"""File readers for Stage 4 TXT, DOCX, and text-based PDF support."""

from pathlib import Path


TXT_EXTENSION = ".txt"
DOCX_EXTENSION = ".docx"
PDF_EXTENSION = ".pdf"
SUPPORTED_EXTENSIONS = (TXT_EXTENSION, DOCX_EXTENSION, PDF_EXTENSION)


def _unsupported_extension_error(file_path: str | Path) -> ValueError:
    path = Path(file_path)
    suffix = path.suffix.lower() or "<none>"
    return ValueError(
        f"Unsupported file extension for Stage 4: {suffix}. "
        "Only .txt, .docx, and .pdf files are supported."
    )


def _ensure_txt_path(file_path: str | Path) -> Path:
    path = Path(file_path)
    if path.suffix.lower() != TXT_EXTENSION:
        suffix = path.suffix or "<none>"
        raise ValueError(
            f"Unsupported file extension for read_txt_file: {suffix}. "
            "Only .txt files are supported by read_txt_file."
        )
    return path


def _ensure_docx_path(file_path: str | Path) -> Path:
    path = Path(file_path)
    if path.suffix.lower() != DOCX_EXTENSION:
        suffix = path.suffix or "<none>"
        raise ValueError(
            f"Unsupported file extension for read_docx_file: {suffix}. "
            "Only .docx files are supported by read_docx_file."
        )
    return path


def _ensure_pdf_path(file_path: str | Path) -> Path:
    path = Path(file_path)
    if path.suffix.lower() != PDF_EXTENSION:
        suffix = path.suffix or "<none>"
        raise ValueError(
            f"Unsupported file extension for read_pdf_file: {suffix}. "
            "Only .pdf files are supported by read_pdf_file."
        )
    return path


def read_txt_file(file_path: str | Path) -> str:
    """Read a UTF-8 TXT file and return its text content."""
    path = _ensure_txt_path(file_path)
    return path.read_text(encoding="utf-8")


def _load_document_class():
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "DOCX support requires python-docx. "
            "Install dependencies from requirements.txt."
        ) from exc

    return Document


def _load_pdf_reader_class():
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "PDF support requires pypdf. "
            "Install dependencies from requirements.txt."
        ) from exc

    return PdfReader


def _iter_docx_text_parts(document) -> list[str]:
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        text_parts.append(paragraph.text)

    return text_parts


def read_docx_file(file_path: str | Path) -> str:
    """Read basic paragraph and simple table text from a local DOCX file."""
    path = _ensure_docx_path(file_path)
    Document = _load_document_class()
    document = Document(path)
    return "\n".join(_iter_docx_text_parts(document))


def read_pdf_file(file_path: str | Path) -> str:
    """Read extractable text from a local text-based PDF file."""
    path = _ensure_pdf_path(file_path)
    PdfReader = _load_pdf_reader_class()
    reader = PdfReader(path)
    text_parts: list[str] = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    text = "\n".join(text_parts)
    if not text.strip():
        raise ValueError(
            "PDF contains no extractable text. Only text-based PDFs are "
            "supported; scanned PDFs and OCR are not supported."
        )

    return text


def extract_text(file_path: str | Path) -> str:
    """Extract text from a supported Stage 4 input file."""
    path = Path(file_path)

    if path.suffix.lower() == TXT_EXTENSION:
        return read_txt_file(path)
    if path.suffix.lower() == DOCX_EXTENSION:
        return read_docx_file(path)
    if path.suffix.lower() == PDF_EXTENSION:
        return read_pdf_file(path)

    raise _unsupported_extension_error(path)
