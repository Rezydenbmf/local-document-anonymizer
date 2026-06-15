"""File writers for Stage 6 TXT, DOCX, PDF-to-TXT, and report support."""

from collections.abc import Callable
from pathlib import Path


TXT_EXTENSION = ".txt"
DOCX_EXTENSION = ".docx"
PDF_EXTENSION = ".pdf"
ANON_SUFFIX = "_ANON"
REPORT_SUFFIX = "_RAPORT"
BATCH_SUMMARY_FILENAME = "_BATCH_SUMMARY.txt"
AnonymizeFunction = Callable[[str], tuple[str, dict[str, int]]]


def _unsupported_extension_error(file_path: str | Path) -> ValueError:
    path = Path(file_path)
    suffix = path.suffix.lower() or "<none>"
    return ValueError(
        f"Unsupported file extension: {suffix}. "
        "Only .txt, .docx, and .pdf files are supported."
    )


def _ensure_txt_path(file_path: str | Path) -> Path:
    path = Path(file_path)
    if path.suffix.lower() != TXT_EXTENSION:
        suffix = path.suffix or "<none>"
        raise ValueError(
            f"Unsupported file extension for TXT output: {suffix}. "
            "Only .txt files are supported by TXT output helpers."
        )
    return path


def _ensure_docx_path(file_path: str | Path) -> Path:
    path = Path(file_path)
    if path.suffix.lower() != DOCX_EXTENSION:
        suffix = path.suffix or "<none>"
        raise ValueError(
            f"Unsupported file extension for DOCX output: {suffix}. "
            "Only .docx files are supported by DOCX output helpers."
        )
    return path


def _ensure_pdf_path(file_path: str | Path) -> Path:
    path = Path(file_path)
    if path.suffix.lower() != PDF_EXTENSION:
        suffix = path.suffix or "<none>"
        raise ValueError(
            f"Unsupported file extension for PDF TXT output: {suffix}. "
            "Only .pdf files are supported by PDF TXT output helpers."
        )
    return path


def build_collision_safe_path(candidate_path: str | Path) -> Path:
    """Return candidate path or the next numbered path if it already exists."""
    path = Path(candidate_path)
    if not path.exists():
        return path

    index = 2
    while True:
        next_path = path.with_name(f"{path.stem}_{index}{path.suffix}")
        if not next_path.exists():
            return next_path
        index += 1


def _output_directory(source_path: str | Path, output_dir: str | Path | None) -> Path:
    if output_dir is None:
        return Path(source_path).parent
    return Path(output_dir)


def build_anonymized_txt_path(
    source_path: str | Path, output_dir: str | Path | None = None
) -> Path:
    """Return the anonymized output path for a TXT source file."""
    path = _ensure_txt_path(source_path)
    return _output_directory(path, output_dir) / f"{path.stem}{ANON_SUFFIX}{TXT_EXTENSION}"


def build_anonymized_docx_path(
    source_path: str | Path, output_dir: str | Path | None = None
) -> Path:
    """Return the anonymized output path for a DOCX source file."""
    path = _ensure_docx_path(source_path)
    return _output_directory(path, output_dir) / f"{path.stem}{ANON_SUFFIX}{DOCX_EXTENSION}"


def build_anonymized_pdf_txt_path(
    source_path: str | Path, output_dir: str | Path | None = None
) -> Path:
    """Return the anonymized TXT output path for a PDF source file."""
    path = _ensure_pdf_path(source_path)
    return _output_directory(path, output_dir) / f"{path.stem}{ANON_SUFFIX}{TXT_EXTENSION}"


def build_report_path(
    source_path: str | Path, output_dir: str | Path | None = None
) -> Path:
    """Return the safe report output path for a supported source file."""
    path = Path(source_path)
    if path.suffix.lower() not in (TXT_EXTENSION, DOCX_EXTENSION, PDF_EXTENSION):
        raise _unsupported_extension_error(path)

    return _output_directory(path, output_dir) / f"{path.stem}{REPORT_SUFFIX}{TXT_EXTENSION}"


def build_batch_summary_path(output_dir: str | Path) -> Path:
    """Return the default batch summary path in an output workspace."""
    return Path(output_dir) / BATCH_SUMMARY_FILENAME


def save_anonymized_txt_copy(
    source_path: str | Path,
    anonymized_text: str,
    output_dir: str | Path | None = None,
) -> Path:
    """Write anonymized UTF-8 TXT content without modifying the source file."""
    if not isinstance(anonymized_text, str):
        raise TypeError("anonymized_text must be a string")

    output_path = build_collision_safe_path(
        build_anonymized_txt_path(source_path, output_dir=output_dir)
    )
    output_path.write_text(anonymized_text, encoding="utf-8")
    return output_path


def save_anonymized_pdf_txt_copy(
    source_path: str | Path,
    anonymized_text: str,
    output_dir: str | Path | None = None,
) -> Path:
    """Write anonymized PDF text as UTF-8 TXT without modifying the PDF."""
    if not isinstance(anonymized_text, str):
        raise TypeError("anonymized_text must be a string")

    output_path = build_collision_safe_path(
        build_anonymized_pdf_txt_path(source_path, output_dir=output_dir)
    )
    output_path.write_text(anonymized_text, encoding="utf-8")
    return output_path


def _load_document_class():
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "DOCX support requires python-docx. "
            "Install dependencies from requirements.txt."
        ) from exc

    return Document


def _merge_counters(target: dict[str, int], source: dict[str, int]) -> None:
    for label, count in source.items():
        target[label] = target.get(label, 0) + count


def _iter_docx_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _replace_paragraph_text(paragraph, anonymized_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(anonymized_text)
        return

    paragraph.runs[0].text = anonymized_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _anonymize_docx_paragraph(
    paragraph,
    anonymize: AnonymizeFunction,
    anonymize_run: AnonymizeFunction | None = None,
) -> dict[str, int]:
    original_text = paragraph.text
    if not original_text:
        return {}

    anonymized_text, counters = anonymize(original_text)
    if not counters:
        return {}

    run_anonymizer = anonymize_run or anonymize
    for run in paragraph.runs:
        run_text, _ = run_anonymizer(run.text)
        run.text = run_text

    if paragraph.text != anonymized_text:
        _replace_paragraph_text(paragraph, anonymized_text)

    return counters


def save_anonymized_docx_copy(
    source_path: str | Path,
    anonymize: AnonymizeFunction,
    anonymize_run: AnonymizeFunction | None = None,
    output_dir: str | Path | None = None,
) -> tuple[Path, dict[str, int]]:
    """Anonymize a local DOCX source and save a separate _ANON.docx copy."""
    if not callable(anonymize):
        raise TypeError("anonymize must be callable")
    if anonymize_run is not None and not callable(anonymize_run):
        raise TypeError("anonymize_run must be callable")

    path = _ensure_docx_path(source_path)
    output_path = build_collision_safe_path(
        build_anonymized_docx_path(path, output_dir=output_dir)
    )
    Document = _load_document_class()
    document = Document(path)
    counters: dict[str, int] = {}

    for paragraph in _iter_docx_paragraphs(document):
        _merge_counters(
            counters,
            _anonymize_docx_paragraph(
                paragraph,
                anonymize,
                anonymize_run=anonymize_run,
            ),
        )

    document.save(output_path)
    return output_path, counters


def save_anonymized_copy(
    source_path: str | Path,
    anonymized_text: str,
    output_dir: str | Path | None = None,
) -> str:
    """Save an anonymized copy and return its path as text."""
    path = Path(source_path)

    if path.suffix.lower() == TXT_EXTENSION:
        return str(save_anonymized_txt_copy(path, anonymized_text, output_dir=output_dir))
    if path.suffix.lower() == DOCX_EXTENSION:
        raise ValueError(
            "DOCX output requires save_anonymized_docx_copy(...), which "
            "uses the existing anonymization engine on the source DOCX."
        )
    if path.suffix.lower() == PDF_EXTENSION:
        return str(
            save_anonymized_pdf_txt_copy(path, anonymized_text, output_dir=output_dir)
        )

    raise _unsupported_extension_error(path)
